from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "outputs" / "롤플레잉_영어학습앱_통합_개발명세서.md"
OUTPUT = ROOT / "outputs" / "롤플레잉_영어학습앱_통합_개발명세서.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT


def style_run(run, bold=False, size=10.5, color="111111"):
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph_with_inline_code(doc, text, style=None):
    paragraph = doc.add_paragraph(style=style)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            style_run(run, size=10, color="0B5D4E")
            run.font.name = "Consolas"
        else:
            run = paragraph.add_run(part)
            style_run(run)
    return paragraph


def add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_width(table)
    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            style_run(run, bold=(row_index == 0), size=9.5)
            if row_index == 0:
                set_cell_shading(cell, "EDEDED")
    doc.add_paragraph()


def apply_document_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "16463F", 14, 8),
        ("Heading 2", 13, "16463F", 12, 6),
        ("Heading 3", 11.5, "445B55", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build():
    source = SOURCE.read_text(encoding="utf-8")
    lines = source.splitlines()

    doc = Document()
    apply_document_styles(doc)

    table_buffer = []
    in_code = False
    code_lines = []

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            rows = []
            for line in table_buffer:
                if re.match(r"^\|[-:| ]+\|$", line):
                    continue
                cells = [cell.strip() for cell in line.strip("|").split("|")]
                rows.append(cells)
            add_table(doc, rows)
            table_buffer = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.15)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            for i, line in enumerate(code_lines):
                if i:
                    paragraph.add_run().add_break()
                run = paragraph.add_run(line)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor.from_string("333333")
            code_lines = []

    for line in lines:
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("|"):
            table_buffer.append(line)
            continue

        flush_table()

        if not line.strip():
            continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(line[2:].strip())
            style_run(run, bold=True, size=21, color="111111")
            paragraph.paragraph_format.space_after = Pt(10)
        elif line.startswith("## "):
            add_paragraph_with_inline_code(doc, line[3:].strip(), style="Heading 1")
        elif line.startswith("### "):
            add_paragraph_with_inline_code(doc, line[4:].strip(), style="Heading 2")
        elif line.startswith("- "):
            paragraph = add_paragraph_with_inline_code(doc, line[2:].strip(), style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(3)
        elif re.match(r"^\d+\. ", line):
            paragraph = add_paragraph_with_inline_code(doc, re.sub(r"^\d+\. ", "", line).strip(), style="List Number")
            paragraph.paragraph_format.space_after = Pt(3)
        else:
            add_paragraph_with_inline_code(doc, line.strip())

    flush_table()
    flush_code()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
